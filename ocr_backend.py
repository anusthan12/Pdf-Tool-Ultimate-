# ocr_backend.py
import os
import io
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename

# OCR and PDF processing libraries
import fitz  # PyMuPDF
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
from docx import Document

# --- App Initialization ---
app = Flask(__name__)
CORS(app)  # Allow requests from all origins

# Configure Tesseract path (Render specific)
if os.environ.get('RENDER'):
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# --- Helper Functions ---
def convert_without_ocr(pdf_bytes):
    """Extracts text directly from a PDF and returns a Word document."""
    doc = Document()
    doc.add_heading('Converted Document (Text Extraction)', 0)
    
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            
            if text.strip():  # Only add non-empty text
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("[No text found on this page]")
                
            if page_num < len(pdf_document) - 1:
                doc.add_page_break()
        
        pdf_document.close()
    except Exception as e:
        doc.add_paragraph(f"Error processing PDF: {str(e)}")
    
    return save_doc_to_memory(doc)

def convert_with_ocr(pdf_bytes):
    """Converts PDF pages to images and uses Tesseract OCR to extract text."""
    doc = Document()
    doc.add_heading('Converted Document (OCR)', 0)
    
    try:
        images = convert_from_bytes(pdf_bytes, dpi=200)  # Reduced DPI for faster processing
        for i, image in enumerate(images):
            # Convert to grayscale for better OCR results
            if image.mode != 'L':
                image = image.convert('L')
            
            text = pytesseract.image_to_string(image, config='--psm 6')
            
            if text.strip():
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(f"[No text detected on page {i+1}]")
                
            if i < len(images) - 1:
                doc.add_page_break()
                
    except Exception as e:
        doc.add_paragraph(f"Error during OCR processing: {str(e)}")
    
    return save_doc_to_memory(doc)

def save_doc_to_memory(doc):
    """Saves a docx.Document object to an in-memory bytes buffer."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- API Endpoints ---
@app.route('/', methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({
        "status": "healthy", 
        "message": "OCR Backend is running!",
        "tesseract_available": os.path.exists('/usr/bin/tesseract')
    }), 200

@app.route('/convert-pdf-to-word', methods=['POST'])
def convert_pdf_to_word():
    """Main conversion endpoint."""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        mode = request.form.get('mode', 'no_ocr')

        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.lower().endswith('.pdf'):
            return jsonify({"error": "Please upload a PDF file"}), 400

        # Read file content
        pdf_bytes = file.read()
        
        if len(pdf_bytes) == 0:
            return jsonify({"error": "Empty file uploaded"}), 400

        # Process based on mode
        if mode == 'ocr':
            buffer = convert_with_ocr(pdf_bytes)
            filename = 'converted_ocr.docx'
        else:
            buffer = convert_without_ocr(pdf_bytes)
            filename = 'converted_text.docx'
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        app.logger.error(f"Conversion error: {str(e)}")
        return jsonify({"error": f"Processing failed: {str(e)}"}), 500

# Handle CORS preflight requests
@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        response = jsonify({})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "*")
        response.headers.add('Access-Control-Allow-Methods', "*")
        return response

# --- Error Handlers ---
@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "File too large. Please upload a smaller PDF."}), 413

@app.errorhandler(500)
def internal_server_error(error):
    return jsonify({"error": "Internal server error occurred."}), 500

# --- Configuration ---
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
